# -*- coding: utf8 -*-

import os
from aiogram import Bot, Dispatcher, executor, types
from aiogram.dispatcher.filters import Command
from docx import Document
from Messages import Message_to_user

TOKEN = *TOKEN*
bot = Bot(token=TOKEN)
dp = Dispatcher(bot=bot)
ingame = False  #
document_with_rules = Document()
game_database = []
msg = Message_to_user()

@dp.message_handler(commands=['start', 'allgames', 'help', 'addgame', 'sleep'])
async def commands_handler(message: types.Message, command: Command.CommandObj):
    global ingame
    global game_database
    global msg

    if command.command =='start':
        game_database = get_database()
        ingame = False
        await message.reply(msg.greeting_response(message.from_user.first_name),
                            reply_markup=types.ReplyKeyboardRemove())
    elif command.command == 'allgames':
        for file in game_database:
            game_name = str(file)
            keyboard = types.InlineKeyboardMarkup()
            keyboard.add(types.InlineKeyboardButton(text=game_name, callback_data=game_name))
        await message.answer("Нажмите для выбора игры", reply_markup=keyboard)
    elif command.command == 'help':
        await message.answer(msg.help_response())
    elif command.command == 'addgame':
        await message.answer(msg.addgame_response())
    elif command.command == 'sleep' and message.from_user.id == 966427877:
        await message.answer(f'Твой ID={message.from_user.id}')
        #дописать массовое сообщение о выключении
    else:
        await message.answer('Такая команда мне не знакома!')


@dp.callback_query_handler()  # обработчик кнопки для списка всех игр
async def answer_to_button(call: types.CallbackQuery):
    global ingame
    global document_with_rules
    if ingame == False:  # логика, отвечающая за поиск игры в базе, активирует кнопки конкретной игры
        request_answer = get_rules(call.data)
        if isinstance(request_answer, str):
            await call.message.answer(request_answer)
        else:
            document_with_rules = request_answer  # работаем с документом (правилами)
            newkeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
            for row in document_with_rules.tables[0].rows:  # делаем кнопки
                button = types.KeyboardButton(text=row.cells[0].text.strip())
                newkeyboard.add(button)
            ingame = True
            await call.message.answer('Вот, смотри! Жми на нужную тебе часть правил в появившейся клавиатуре и всё получится!', reply_markup=newkeyboard)



    #await call.message.answer(call.data)

# Функция реакции на сообщения
@dp.message_handler(content_types=["text"])
async def answer_to_user(message: types.Message):
    global ingame
    global document_with_rules
    if ingame == False:  # логика, отвечающая за поиск игры в базе, активирует кнопки конкретной игры
        request_answer = get_rules(message.text)
        if isinstance(request_answer, str):
            await message.answer(request_answer)
        else:
            document_with_rules = request_answer  # работаем с документом (правилами)
            newkeyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
            for row in document_with_rules.tables[0].rows:  # делаем кнопки
                button = types.KeyboardButton(text=row.cells[0].text.strip())
                newkeyboard.add(button)
            ingame = True
            await message.answer('Вот, смотри\nЧтобы выйти из этого меню нажми кнопку start в меню слева', reply_markup=newkeyboard)

# соответствие кнопок вариантам ответов
    elif ingame == True:
        # поставить запрет на печать с клавиатуры хз можно ли
        button_rule = []
        answer_rule = []
        for i in range(0, len(document_with_rules.tables[0].rows)):
            button_rule.append(document_with_rules.tables[0].rows[i].cells[0].text.strip())
            answer_rule.append(document_with_rules.tables[0].rows[i].cells[1].text.strip())

        for i in range(0, len(button_rule)):
            if message.text == button_rule[i]:
                await message.answer(answer_rule[i])
            else:
                pass


# функция должна возвращать корректные имена игры + сделать на множеств. число игр и выдачу инлайн предложений
def get_game_name(user_request):

    global game_database
    answer = ('Такой игры я еще не знаю! '
                '\nЕсли хочешь, то можешь посмотреть в раздел о добавлении игр в меню!')
    cor_name = ''

    for games in game_database:
        if user_request.lower() in games.lower():
            cor_name = games
        else:
            pass

    if cor_name in game_database:
        answer = f'Есть похожая игра. Её корректное название - {cor_name}'
    else:
        pass

    return answer


# непосредственно выдает текстовый файл с правилами или отправляет искать верное имя
def get_rules(user_request):
    global game_database
    correct_name = ''

    for games in game_database:
        if user_request.lower() == games.lower():
            correct_name = games
        else:
            pass

    if correct_name != '':
        path = (f'C:/Users/Squeed/Desktop/LazyRules/{user_request}.docx');
        return (Document(path))
    else:
        return get_game_name(user_request)


# получение базы данных
def get_database():
    """Получение базы данных"""
    files = os.listdir('C:/Users/Squeed/Desktop/LazyRules/')
    database = []
    for file in files:
        game_name = str(file)[:-5:]
        database.append(game_name)
    return database


if __name__ == '__main__':
    executor.start_polling(dp)