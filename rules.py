# -*- coding: utf8 -*-
import os
from docx import Document

def get_database():
    """Получение базы данных"""
    files = os.listdir('C:/Users/Squeed/Desktop/LazyRules/')
    database = []
    for file in files:
        game_name = str(file)[:-5:]
        database.append(game_name)
    db = tuple(database)
    return db

class CurrentRule():
    rule_buttons = []
    rule_answers = []

    def __init__(self, rule_dock=None):
        self.rule_buttons = []
        self.rule_answers = []
        if isinstance(rule_dock, type(Document())):
            for i in range(0, len(rule_dock.tables[0].rows)):
                self.rule_buttons.append(rule_dock.tables[0].rows[i].cells[0].text.strip())
            for i in range(0, len(rule_dock.tables[0].rows)):
                self.rule_answers.append(rule_dock.tables[0].rows[i].cells[1].text.strip())
        else:pass



class GameRules():
    _game_list = get_database()

    def get_game(self, user_request):
        answer = ('Такой игры я еще не знаю! '
                  '\nЕсли хочешь, то можешь посмотреть в раздел о добавлении игр в меню!')
        for games in self._game_list:
            if user_request.lower() in games.lower():
                correct_name = games
                if user_request.lower() == games.lower():
                    path = ('C:/Users/Squeed/Desktop/LazyRules/{}.docx').format(user_request);
                    return (Document(path))
                elif correct_name in self._game_list:
                    answer = 'Есть игра с похожим названием. Её корректное название - {}'.format(correct_name)
            else:
                pass
        return answer